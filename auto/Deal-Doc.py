# -*- coding: utf-8 -*-
# @Author  : Zoey Zhao
# @Time    : 2020/11/30 16:21
# @Function: 自动处理Word
import docx
from docx.shared import Cm

# 检查内容
checkItem = (
    'A1',
    'A2',
)

# 替换内容
changeItem = {
    'A1': '1234',# 流动资产
    'A2': '4321'# 货币资金
}



doc = docx.Document('result.docx')

tables = doc.tables

## 遍历全部单元格（下标方式）
tableLen = len(doc.tables)
#print("表格数量", 表格数量)
for tableID in range(0, tableLen):
    #doc.tables[表格编号]                                         ## 遍历每一个表格
    rows = len(doc.tables[tableID].rows)
    #print("行数", 行数)
    for rowsID in range(0, rows):
        #doc.tables[表格编号].rows[行编号]                        ## 遍历每一个表格的每一行
        cols = len(doc.tables[tableID].rows[rowsID].cells)
        #print("列数", 列数)
        for colsID in range(0, cols):
            #doc.tables[表格编号].rows[行编号].cells[列编号]      ## 遍历每一个表格的每一行的每一列
            tmp = doc.tables[tableID].rows[rowsID].cells[colsID].text
            # print(tmp)
            if tmp in checkItem:
                print(doc.tables[tableID].rows[rowsID].cells[colsID].text)
                print(changeItem[tmp])
                doc.tables[tableID].rows[rowsID].cells[colsID].text = changeItem[tmp]
                print(doc.tables[tableID].rows[rowsID].cells[colsID].text)




# for i in range(len(tables)):
#     print(tables[i])
doc.save('20201130.docx')
